# -*- coding: UTF-8 -*-
import re
import sys
import os
import time
import csv
import argparse
import codecs
from docx import Document

import Spider


def download_case(wenshu, round):
    if round == 2:
        path = 'Download/'
        col_name = 'name'
        col_id = 'name'
        col_date = 'date'
    elif round == 1:
        path = 'Download1/'
        col_name = 'name1'
        col_id = 'id1'
        col_date = 'date1'
    else:
        print('invalid round')
    download_list = ['Y'] * len(wenshu.case[col_name])
    print(len(download_list))
    for i in range(len(wenshu.case[col_name])):
        #print(i)
        file_name = path + wenshu.case[col_name][i] + wenshu.case[col_date][i] + '.docx'
        if not os.path.exists(file_name) and wenshu.case[col_name][i][0] != '[' and (wenshu.case[col_name][i] != 'None' and wenshu.case[col_name][i] != 'na'):
            if round == 1:
                search_criteria = '案号:' + wenshu.case['doc_id'][i] + ',审判程序:一审' + ',法院地域:四川省,关键词:离婚'
                wenshu.setSearchCriteria(search_criteria)
                wenshu.setDownloadConditions()
            wenshu.downloadDocument(path,
                                    wenshu.case[col_name][i],
                                    wenshu.case[col_id][i],
                                    wenshu.case[col_date][i])
            docsize = os.path.getsize(file_name)
            # if docsize < 80k, it may corrupt. resend request
            if docsize < 80000:
                print('file %s is invalid' % file_name)
                input("Refresh the Page and Enter:")
                wenshu.downloadDocument(path,
                                        wenshu.case[col_name][i],
                                        wenshu.case[col_id][i],
                                        wenshu.case[col_date][i])
                docsize = os.path.getsize(file_name)
                print('docsize is %s' % docsize)
                download_list[i] = 'N'
            else:
                download_list[i] = 'N'
    if round == 2:
        wenshu.case['download'] = download_list
    elif round == 1:
        wenshu.case['download1'] = download_list
    else:
        print('invalid round')

    

def get_case_info(wenshu):
    wenshu.getTotalItemNumber()
    print("Total case number is %s" % wenshu.total_items)
    wenshu.getCaseList(wenshu.total_items)


def write_2_csv(case):
    with open('case.csv', 'w', newline='', encoding='utf-8') as csvfile:
        csvfile.write(u'\ufeff')
        fieldnames = ['name', 'id', 'exist']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for i in range(len(case['name'])):
            writer.writerow({'name': case['name'][i],
                             'id': case['id1'][i],
                             'exist': case['exist'][i]})


def get_case_1st_id(wenshu):
    doc_id_1st_list = ['None'] * len(wenshu.case['name'])
    for i in range(len(wenshu.case['name'])):
        file_name = 'Download/' + wenshu.case['name'][i] + wenshu.case['date'][i] + '.docx'    
        if wenshu.case['download'][i] == 'Y':
            print('Processing document %s %s '%(i, wenshu.case['name'][i]))
            doc = read_doc(file_name)
            doc_id_1st_list[i] = process_doc_data(doc)
    wenshu.case['doc_id'] = doc_id_1st_list


def read_doc(doc_name=None):
    #   打开文档
    try:
        document = Document(doc_name) if doc_name else sys.exit(0)
    except docx.opc.exceptions.PackageNotFoundError:
        print("Document %s is invalid" % doc_name)
        s = 'NA'
        return s
    #   读取每段资料
    l = [paragraph.text for paragraph in document.paragraphs]
    s = ''.join(str(e) for e in l)

    return s


def process_doc_data(doc_data=None):
    id_1st = re.search('.\d\d\d\d.\w+?民.?初.+?号', doc_data)
    if id_1st:
        return id_1st.group()
    else:
        return 'None'
        
        
def search(wenshu, wenshu1):    
    for i in range(len(wenshu.case['name'])):
    #for i in range(0, 2):
        print(i)
        
        if wenshu.case['doc_id'][i] == 'None':
            print('1st case id %s is not exist, Skip %s' % (wenshu.case['doc_id'][i], wenshu.case['name'][i]))
        elif wenshu.case['name1'][i] == 'na':
            print("Search document with id1 %s" % wenshu.case['doc_id'][i])
            search_criteria = '案号:' + wenshu.case['doc_id'][i] + ',审判程序:一审' + ',法院地域:四川省,关键词:离婚'
            wenshu1.setSearchCriteria(search_criteria)
            get_case_info(wenshu1)
            #print(wenshu1.case)
            if len(wenshu1.case['name']) == 1:
                #print('len(wenshu1.case[\'name\']) %s'% len(wenshu1.case['name']))
                wenshu.case['name1'][i] = wenshu1.case['name'][0]
                wenshu.case['id1'][i] = wenshu1.case['id'][0]
                wenshu.case['date1'][i] = wenshu1.case['date'][0]
                wenshu.case['case_id1'][i] = wenshu1.case['case_id'][0]
                if wenshu1.case['case_id'][0] == wenshu.case['doc_id'][i]:
                    wenshu.case['match'][i] = 'Y'
            elif len(wenshu1.case['name']) == 0:
                #print('len(wenshu1.case[\'name\']) %s'% len(wenshu1.case['name']))
                wenshu.case['name1'][i] = 'None'
                wenshu.case['id1'][i] = 'None'
                wenshu.case['date1'][i] = 'None'
                wenshu.case['case_id1'][i] = 'None'
            else:
                wenshu.case['name1'][i] = wenshu1.case['name']
                wenshu.case['id1'][i] = wenshu1.case['id']
                wenshu.case['date1'][i] = wenshu1.case['date']
                wenshu.case['case_id1'][i] = wenshu1.case['case_id']
        else:
            print('Doc %s exist, skip' % wenshu.case['name1'][i])
        dump2csv(wenshu,'phase4')


def dump2csv(wenshu, surfix):
    with open('case' + surfix + '.csv', 'w', newline='', encoding='utf-8_sig') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(wenshu.case.keys())
        writer.writerows(zip(*wenshu.case.values()))


def read_csv(wenshu, surfix):
    with open('case' + surfix + '.csv', encoding='utf-8_sig') as csvfile:
        reader = csv.DictReader(csvfile)
        case = dict.fromkeys(reader.fieldnames)
        for key in case:
            case[key] = []
        #print(case)
        for row in reader:
            for key in case:
                case[key].append(row[key])
    wenshu.case = case

def clean_data(wenshu):
    multi_count = 0
    none_count = 0
    empty_count = 0
    for i in range(len(wenshu.case['name1'])):
        if not wenshu.case['name1'][i]:
            empty_count += 1
        elif wenshu.case['name1'][i] == 'None':
            none_count += 1
        else:
            if wenshu.case['name1'][i][0] == '[':
                print(wenshu.case['name1'][i])
                multi_count += 1
    print('empty_count is %s, none_count is %s, multi_count is %s' % (empty_count, none_count, multi_count))
    
# Phase 1: Search and get 2nd case list, download all of them,
#          dump case name list into a csv file.
def phase1(wenshu):
    # Get 2nd case list with search criteria.
    get_case_info(wenshu)
    dump2csv(wenshu, 'phase1')
    

def phase2(wenshu):
    # Read csv file and get case list.
    read_csv(wenshu, 'phase1')
    download_case(wenshu, 2)
    dump2csv(wenshu,'phase2')

def phase3(wenshu):
    read_csv(wenshu, 'phase2')
    get_case_1st_id(wenshu)
    dump2csv(wenshu,'phase3')
    
    
def phase4(wenshu, wenshu1):
    if not os.path.exists('casephase4.csv'):
        read_csv(wenshu, 'phase3')
        wenshu.case['name1'] = ['na'] * len(wenshu.case['name'])
        wenshu.case['id1'] = ['na'] * len(wenshu.case['name'])
        wenshu.case['date1'] = ['na'] * len(wenshu.case['name'])
        wenshu.case['case_id1'] = ['na'] * len(wenshu.case['name'])
        wenshu.case['match'] = ['N'] * len(wenshu.case['name'])
        dump2csv(wenshu,'phase4')
    read_csv(wenshu, 'phase4')    
    search(wenshu, wenshu1)
    dump2csv(wenshu,'phase4')

def phase5(wenshu):
    read_csv(wenshu, 'phase4')
    clean_data(wenshu)
    download_case(wenshu, 1)
    dump2csv(wenshu,'phase5')
    
    
def main():
    desc = "Select a phase to run"
    parser = argparse.ArgumentParser(description=desc)
    parser.add_argument('-p', '--phase', action='store')
    args = parser.parse_args()
    search_criteria = "案件类型:民事案件,法院地域:四川省,四级案由:离婚纠纷,审判程序:二审"
    wenshu = Spider.WenShu()
    wenshu.setSearchCriteria(search_criteria)
    # Phase 1: Search 2nd case and document them into a csv file.
    # Phase 2: Read case list from csv file and download all of them.
    # Phase 3: Analyse 2nd case list and get 1st case id.
    # Phase 4: Search 1st case.
    # Phase 5: Download 1st case.
    if args.phase == 'all':
        print('phase 1')
        print('phase 2')
        print('phase 3')
    elif args.phase == '1':
        print('phase 1')
        phase1(wenshu)
    elif args.phase == '2':
        print('phase 2')
        phase2(wenshu)
    elif args.phase == '3':
        print('phase 3')
        phase3(wenshu)
    elif args.phase == '4':
        print('phase 4')
        wenshu1 = Spider.WenShu()
        phase4(wenshu, wenshu1)
    elif args.phase == '5':
        phase5(wenshu)
    else:
        print('invalid')

    sys.exit(0)


# Debug doc regression
#    file_name = 'Download/陈某与熊某离婚纠纷二审民事判决书2016-06-27.docx'
#    doc = read_doc(file_name)
#    print(doc)
#    id_1st = re.search('.\d\d\d\d.\w+?民.?初.+?号', doc)
#    print(id_1st)     

if __name__ == "__main__":
    main()
